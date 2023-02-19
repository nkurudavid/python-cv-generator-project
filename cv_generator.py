from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)

document = Document()

# profile picture
document.add_picture(
    "my_image.jpg",
    width=Inches(1.5),
)

# names, phone number and email
full_name = input("What is your name?\n>>")
speak("Hello " + full_name + ", how are you today?")

speak("Enter your phone number")
phone_number = input("Enter your phone number?\n>>")

speak("Enter your email address")
email = input("Enter your email address?\n>>")

document.add_paragraph(
    full_name + "\n" + phone_number + "\n" + email
)

# about me
document.add_heading("About me")

speak("Tell me about yourself...")
document.add_paragraph(
    input("Tell me about yourself...\n>>")
)

# work experience
document.add_heading("Work Experience")
p = document.add_paragraph()

speak("You need to specify your work experience")
company = input("Your work experience\nEnter company : ")

speak(f"When have you started working at {company}?")
from_date = input("From Date : ")

speak(f"When was your last time working at {company}?")
to_date = input("To Date : ")

p.add_run(company + " ").bold = True
p.add_run(from_date + " " + to_date + "\n").italic = True

speak(f"Describe your experience at {company}")
experience_details = input("Describe your experience at " + company)
p.add_run(experience_details)

# More experiences
while True:
    speak("Do you have more experiences?")
    has_more_experiences = input("Do you have more experiences? Yes or No\n>>")
    
    if has_more_experiences.lower() == "yes":
        p = document.add_paragraph()
        company = input("Enter company : ")
        
        speak(f"When have you started working at {company}?")
        from_date = input("From Date : ")
        
        speak(f"When was your last time working at {company}?")
        to_date = input("To Date : ")

        p.add_run(company + " ").bold = True
        p.add_run(from_date + " - " + to_date + "\n").italic = True

        speak(f"Describe your experience at {company}")
        experience_details = input(f"Describe your experience at {company} \n>>")
        p.add_run(experience_details)
        
    else:
        break


# skills
document.add_heading("Skills")

speak("Enter your skill")
skills = input("Enter your skill \n>> ")
p = document.add_paragraph(skills)
p.style = "List Bullet"

# More skills
while True:
    speak("Do you have more skills?")
    has_more_skills = input("Do you have more skills? Yes or No\n>>")
    
    if has_more_skills.lower() == "yes":
        speak("Enter your skill")
        skills = input("Enter your skill \n>> ")
        p = document.add_paragraph(skills)
        p.style = "List Bullet"
        
    else:
        break
    

speak("Thank you so much for your participant, your CV is completely generated.")

# footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV generated using Davbencode in python course."


document.save("cv.docx")